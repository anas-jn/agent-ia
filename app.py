# 📦 Imports
import streamlit as st
import io
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF

# 🧠 LangChain
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.chains.question_answering import load_qa_chain

# ⚙️ Configuration
st.set_page_config(page_title="Portfolio IA - Anas",
                   page_icon="🤖", layout="wide")
if "pitch_text" not in st.session_state:
    st.session_state.pitch_text = ""

# 🎨 Branding
st.sidebar.image("assets/logo.jpg", width=120)
st.sidebar.markdown("## 🧭 Navigation IA")
menu = st.sidebar.radio("📂 Modules disponibles :", [
    "🏠 Accueil - Agent IA Chat",
    "📄 Analyse de CV PDF",
    "⚡ Pitch Express",
    "🧠 Générateur de quiz ",
    "💬 Agent support client",
    "📄 Analyse juridique IA",
    "❓ FAQ intelligente"
    
])
st.sidebar.markdown("---")
st.sidebar.markdown("👤 [Profil GitHub](https://github.com/anasjannaj)")
st.sidebar.markdown("📬 Contact : anas@example.com")

# 🏠 Agent IA Chat
if menu == "🏠 Accueil - Agent IA Chat":
    st.title("💬 Agent IA - Chat Général")
    user_input = st.text_input("🗨️ Votre question")
    if user_input:
        llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0.7)
        response = llm.predict(user_input)
        st.markdown("### 🤖 Réponse")
        st.success(response)

# 📄 Analyse de CV PDF
elif menu == "📄 Analyse de CV PDF":
    st.title("📄 Analyse de CV PDF")
    pdf = st.file_uploader("📄 Fichier PDF", type="pdf")
    if pdf:
        reader = PdfReader(pdf)
        raw_text = "".join([page.extract_text() for page in reader.pages])
        splitter = CharacterTextSplitter(
            separator="\n", chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_text(raw_text)
        embeddings = OpenAIEmbeddings()
        db = FAISS.from_texts(chunks, embedding=embeddings)

        st.markdown("### 💡 Suggestions rapides")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📌 Compétences principales"):
                st.session_state["question"] = "Quelles sont les compétences principales du candidat ?"
            if st.button("🧠 Expérience pertinente"):
                st.session_state["question"] = "Quelle est l'expérience la plus pertinente ?"
            if st.button("🎯 Matching IA"):
                st.session_state["question"] = "Ce profil correspond-il à un poste de développeur IA ?"
        with col2:
            if st.button("📈 Points forts & faibles"):
                st.session_state["question"] = "Quels sont les points forts et axes d'amélioration ?"
            if st.button("📝 Résumé professionnel"):
                st.session_state["question"] = "Génère un résumé professionnel du CV."

        question = st.text_input(
            "❓ Posez votre question ici", value=st.session_state.get("question", ""))
        if question:
            docs = db.similarity_search(question)
            llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0)
            chain = load_qa_chain(llm, chain_type="stuff")
            response = chain.run(input_documents=docs, question=question)
            st.markdown("### 🧠 Réponse IA")
            st.success(response)

            format = st.selectbox("📁 Choisissez le format de téléchargement :", [
                                  "TXT", "DOCX", "PDF"])
            if format == "TXT":
                st.download_button("📥 Télécharger en TXT",
                                   data=response, file_name="reponse.txt")
            elif format == "DOCX":
                doc = Document()
                doc.add_paragraph(response)
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button("📥 Télécharger en DOCX",
                                   data=buffer, file_name="reponse.docx")
            elif format == "PDF":
                pdf_file = FPDF()
                pdf_file.add_page()
                pdf_file.set_font("Arial", size=12)
                for line in response.split('\n'):
                    pdf_file.cell(200, 10, txt=line, ln=True)
                pdf_output = io.BytesIO(
                    pdf_file.output(dest='S').encode('latin-1'))
                st.download_button("📥 Télécharger en PDF",
                                   data=pdf_output, file_name="reponse.pdf")
    else:
        st.info("📄 Téléverse un fichier PDF pour commencer.")

# ⚡ Pitch Express
elif menu == "⚡ Pitch Express":
    st.title("⚡ Pitch Express")
    pdf = st.file_uploader("📄 Fichier PDF", type="pdf")
    st.markdown("Résumé rapide de ton profil en 30 secondes.")
    if pdf and st.button("🎤 Générer le pitch"):
        reader = PdfReader(pdf)
        raw_text = "".join([page.extract_text() for page in reader.pages])
        prompt = f"Fais un résumé professionnel et percutant du profil suivant :\n{raw_text}\nRésumé en 30 secondes."
        llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0.5)
        pitch = llm.predict(prompt)
        st.session_state.pitch_text = pitch

    if st.session_state.pitch_text:
        st.markdown("### 🎯 Résumé Express")
        st.info(st.session_state.pitch_text)
        format = st.selectbox("📁 Format de téléchargement :", [
                              "TXT", "DOCX", "PDF"])
        if format == "TXT":
            st.download_button("📥 Télécharger le pitch (TXT)",
                               data=st.session_state.pitch_text, file_name="pitch_express.txt")
        elif format == "DOCX":
            doc = Document()
            doc.add_paragraph(st.session_state.pitch_text)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("📥 Télécharger le pitch (DOCX)",
                               data=buffer, file_name="pitch_express.docx")
        elif format == "PDF":
            pdf_file = FPDF()
            pdf_file.add_page()
            pdf_file.set_font("Arial", size=12)
            for line in st.session_state.pitch_text.split('\n'):
                pdf_file.cell(200, 10, txt=line, ln=True)
            pdf_output = io.BytesIO(
                pdf_file.output(dest='S').encode('latin-1'))
            st.download_button("📥 Télécharger le pitch (PDF)",
                               data=pdf_output, file_name="pitch_express.pdf")
    elif not pdf:
        st.info("📄 Téléverse un fichier PDF pour générer ton pitch.")


# 🧠 Générateur de quiz IA
elif menu == "🧠 Générateur de quiz ":
    st.title("🧠 Générateur de quiz IA")
    pdf = st.file_uploader("📄 Fichier source", type="pdf")

    if pdf and st.button("🎓 Générer des questions"):
        reader = PdfReader(pdf)
        raw_text = "".join([page.extract_text() for page in reader.pages])
        prompt = f"Génère 5 questions à choix multiples sur ce contenu :\n{raw_text}"
        llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0.7)
        quiz = llm.predict(prompt)
        st.session_state.quiz_text = quiz

    if "quiz_text" in st.session_state and st.session_state.quiz_text:
        st.markdown("### 📋 Quiz généré")
        st.write(st.session_state.quiz_text)

        format = st.selectbox("📁 Format de téléchargement :", [
                              "TXT", "DOCX", "PDF"])
        if format == "TXT":
            st.download_button("📥 Télécharger le quiz (TXT)",
                               data=st.session_state.quiz_text, file_name="quiz_ia.txt")
        elif format == "DOCX":
            doc = Document()
            doc.add_paragraph(st.session_state.quiz_text)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("📥 Télécharger le quiz (DOCX)",
                               data=buffer, file_name="quiz_ia.docx")
        elif format == "PDF":
            pdf_file = FPDF()
            pdf_file.add_page()
            pdf_file.set_font("Arial", size=12)
            for line in st.session_state.quiz_text.split('\n'):
                pdf_file.cell(200, 10, txt=line, ln=True)
            pdf_output = io.BytesIO(
                pdf_file.output(dest='S').encode('latin-1'))
            st.download_button("📥 Télécharger le quiz (PDF)",
                               data=pdf_output, file_name="quiz_ia.pdf")


# 💬 Agent support client


elif menu == "💬 Agent support client":
    st.title("💬 Agent IA - Support Client")
    context = st.text_area("📝 Contexte produit/service",
                           "Notre entreprise vend des logiciels de gestion pour PME...")
    question = st.text_input("❓ Question du client")

    if question:
        prompt = f"Voici le contexte produit/service : {context}\nRéponds à la question suivante : {question}"
        llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0.7)
        response = llm.predict(prompt)
        st.markdown("### 🤖 Réponse IA")
        st.success(response)

        format = st.selectbox("📁 Format de téléchargement :", [
                              "TXT", "DOCX", "PDF"])
        if format == "TXT":
            st.download_button("📥 Télécharger la réponse (TXT)",
                               data=response, file_name="reponse_support.txt")
        elif format == "DOCX":
            doc = Document()
            doc.add_paragraph(response)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("📥 Télécharger la réponse (DOCX)",
                               data=buffer, file_name="reponse_support.docx")
        elif format == "PDF":
            pdf_file = FPDF()
            pdf_file.add_page()
            pdf_file.set_font("Arial", size=12)
            for line in response.split('\n'):
                pdf_file.cell(200, 10, txt=line, ln=True)
            pdf_output = io.BytesIO(
                pdf_file.output(dest='S').encode('latin-1'))
            st.download_button("📥 Télécharger la réponse (PDF)",
                               data=pdf_output, file_name="reponse_support.pdf")


# ❓ FAQ intelligente
elif menu == "❓ FAQ intelligente":
    st.title("❓ FAQ intelligente")
    faq_input = st.text_input("💬 Pose ta question sur le profil d’Anas")
    if faq_input:
        context = """
        Anas Jannaj est développeur digital, spécialisé en Python, Flutter, IA, Streamlit, Supabase.
        Il a créé jimla.ma, des LMS, des générateurs de quiz IA.
        Il vise une carrière freelance avec des outils locaux et puissants.
        Il est basé au Maroc, parle français et anglais, et est très actif dans l’apprentissage rapide.
        """
        prompt = f"Voici le contexte : {context}\nRéponds à cette question : {faq_input}"
        llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0.5)
        response = llm.predict(prompt)
        st.markdown("### 🤖 Réponse IA")
        st.success(response)
